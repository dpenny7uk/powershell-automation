"""
Alteryx Migration - SLT/ITLT Briefing Document
================================================
Generates a Word document summarising the key findings from the
scoped migration report for senior leadership.

Usage:
    python Build-SLTBriefing.py <scoped_migration_report.xlsx> [output.docx]
"""

import sys, os
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from openpyxl import load_workbook


def load_data(path):
    """Load the scoped migration report."""
    wb = load_workbook(path, data_only=True)

    def sheet_to_dicts(name):
        ws = wb[name]
        headers = [cell.value for cell in ws[1]]
        rows = []
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
            rows.append(dict(zip(headers, row)))
        return rows

    return {
        "workflows": sheet_to_dicts("Active Workflows"),
        "schedules": sheet_to_dicts("Schedules"),
        "jobs": sheet_to_dicts("Job Activity") if "Job Activity" in wb.sheetnames else [],
    }


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def build_document(data, output_path):
    """Generate the SLT briefing document."""
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    wf = data["workflows"]
    scheds = data["schedules"]
    jobs = data["jobs"]

    # ── Compute all stats ────────────────────────────────────────────────
    total_gallery = 727  # from the summary sheet
    total_active = len(wf)
    total_inactive = total_gallery - total_active

    status_counts = Counter(r.get("Activity Status") for r in wf)
    scheduled = status_counts.get("Scheduled", 0)
    recent = status_counts.get("Active (Unscheduled - Recent)", 0)
    historic = status_counts.get("Active (Unscheduled - Historic)", 0)

    tier_counts = Counter(r.get("Tier") for r in wf)
    t1 = tier_counts.get("Tier 1 - Simple ETL", 0)
    t2 = tier_counts.get("Tier 2 - Transform & Orchestrate", 0)
    t3 = tier_counts.get("Tier 3 - Predictive/Spatial/Code", 0)
    t4 = tier_counts.get("Tier 4 - Self-Service/Interface", 0)

    has_owner = sum(1 for r in wf if r.get("Owner") and str(r["Owner"]).strip())
    no_owner = total_active - has_owner

    source_counts = Counter(r.get("Ownership Source", "") for r in wf)

    enabled_scheds = sum(1 for s in scheds if s.get("enabled") == True)
    disabled_scheds = sum(1 for s in scheds if s.get("enabled") == False)
    errored_scheds = sum(1 for s in scheds if s.get("state") == "Errored")

    has_macros = sum(1 for r in wf if r.get("Macros") and str(r["Macros"]).strip())
    has_conns = sum(1 for r in wf if r.get("Data Connections") and str(r["Data Connections"]).strip())

    ad_hoc_users = set()
    total_ad_hoc = 0
    for j in jobs:
        creator = j.get("Created By", "")
        if creator and str(creator) not in ("", "None", "Unknown"):
            ad_hoc_users.add(creator)
        total_ad_hoc += (j.get("Run Count") or 0)

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_heading("Alteryx Server Migration", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Scoping Assessment & Key Findings")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("April 2026 | Data & Analytics")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # ── Executive Summary ────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        f"An automated audit of the Alteryx Server Gallery has identified "
        f"{total_gallery} workflows in total. Of these, {total_inactive} "
        f"({total_inactive * 100 // total_gallery}%) are inactive and can be "
        f"decommissioned immediately. The remaining {total_active} workflows "
        f"are actively used and form the migration scope."
    )

    doc.add_paragraph(
        f"Of the {total_active} active workflows, {scheduled} run on automated "
        f"schedules and are the highest priority for migration. A further "
        f"{recent} have been run manually since January 2025, and {historic} "
        f"have historical activity but no recent usage."
    )

    doc.add_paragraph(
        f"No Tier 3 (predictive/spatial) workflows were found, which removes "
        f"the most complex migration category. However, {t4} workflows are "
        f"Tier 4 (self-service analytic apps with user-facing interfaces) and "
        f"require a product decision on what replaces the Alteryx Gallery UI."
    )

    # ── Key Numbers ──────────────────────────────────────────────────────
    doc.add_heading("Key Numbers", level=1)

    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Total workflows in gallery", total_gallery],
            ["Inactive (decommission candidates)", f"{total_inactive} ({total_inactive * 100 // total_gallery}%)"],
            ["Active (in scope for migration)", total_active],
            ["Scheduled (must migrate)", scheduled],
            ["Unscheduled - recent (ran since Jan 2025)", recent],
            ["Unscheduled - historic only", historic],
        ],
        col_widths=[10, 6],
    )

    doc.add_paragraph()

    # ── Migration Complexity ─────────────────────────────────────────────
    doc.add_heading("Migration Complexity", level=1)

    doc.add_paragraph(
        "Workflows have been classified into tiers based on the tools and "
        "patterns they use. This determines which target platform is appropriate "
        "and the level of effort required."
    )

    add_styled_table(doc,
        ["Tier", "Count", "%", "Target Platform"],
        [
            ["Tier 1 - Simple ETL", t1, f"{t1 * 100 // total_active}%", "Any: ADF, Python, Power BI Dataflows, Databricks"],
            ["Tier 2 - Transform & Orchestrate", t2, f"{t2 * 100 // total_active}%", "ADF or Databricks (orchestration needed)"],
            ["Tier 3 - Predictive/Spatial/Code", t3, f"{t3 * 100 // total_active}%" if total_active else "0%", "Databricks or standalone Python"],
            ["Tier 4 - Self-Service/Interface", t4, f"{t4 * 100 // total_active}%", "Needs product decision (user-facing UI)"],
        ],
        col_widths=[5.5, 2, 1.5, 8],
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Complexity indicators: ")
    run.bold = True
    p.add_run(
        f"{has_macros} workflows ({has_macros * 100 // total_active}%) depend on shared macros, "
        f"and {has_conns} ({has_conns * 100 // total_active}%) use named database connections "
        f"that will need remapping. No embedded SQL was found."
    )

    # ── Ownership ────────────────────────────────────────────────────────
    doc.add_heading("Ownership", level=1)

    doc.add_paragraph(
        f"Ownership has been resolved for {has_owner} of {total_active} workflows "
        f"({has_owner * 100 // total_active}%) through a combination of schedule "
        f"ownership data, job execution history, and workflow naming patterns."
    )

    add_styled_table(doc,
        ["Resolution Method", "Count"],
        [
            ["Schedule Owner (from Alteryx API)", source_counts.get("Schedule Owner", 0)],
            ["Job History (last person to run it)", source_counts.get("Job History", 0)],
            ["Name Pattern (inferred, unconfirmed)", source_counts.get("Name Pattern (unconfirmed)", 0)],
            ["Team Inference (team known, no contact)", source_counts.get("Team Inference (unconfirmed)", 0)],
            ["Unresolved", no_owner],
        ],
        col_widths=[10, 3],
    )

    doc.add_paragraph()

    if no_owner > 0:
        doc.add_paragraph(
            f"The {no_owner} unresolved workflows should be investigated via "
            f"the Alteryx Server MongoDB audit logs or escalated to team leads."
        )

    # ── Schedule Health ──────────────────────────────────────────────────
    doc.add_heading("Schedule Health", level=1)

    doc.add_paragraph(
        f"There are {len(scheds)} schedules on the server. Only {enabled_scheds} "
        f"are enabled and active. {disabled_scheds} are disabled and represent "
        f"a cleanup opportunity."
    )

    if errored_scheds:
        p = doc.add_paragraph()
        run = p.add_run(f"{errored_scheds} schedules are in an errored state. ")
        run.bold = True
        p.add_run(
            "These workflows may already be broken and should be investigated "
            "before migration planning begins. Some may be decommission candidates."
        )

    # ── Ad-Hoc Usage ─────────────────────────────────────────────────────
    if jobs:
        doc.add_heading("Ad-Hoc Usage", level=1)

        doc.add_paragraph(
            f"{len(jobs)} unscheduled workflows have been run manually since "
            f"January 2025, totalling {total_ad_hoc} runs across "
            f"{len(ad_hoc_users)} unique users. "
        )

        # Top 5 most-run workflows
        sorted_jobs = sorted(jobs, key=lambda r: r.get("Run Count", 0) or 0, reverse=True)
        top_5 = sorted_jobs[:5]
        add_styled_table(doc,
            ["Workflow", "Runs", "User"],
            [[j.get("Workflow Name", ""), j.get("Run Count", 0), j.get("Created By", "")] for j in top_5],
            col_widths=[8, 2, 6],
        )

        doc.add_paragraph()

        doc.add_paragraph(
            "Workflows with high manual run counts may benefit from scheduling "
            "or replacement with a self-service solution on the target platform."
        )

    # ── Risk & Timeline ──────────────────────────────────────────────────
    doc.add_heading("Risks to September 2026 Deadline", level=1)

    risks = [
        [
            "Tier 4 product decision",
            "High",
            f"{t4} workflows have user-facing interfaces. No replacement has "
            "been agreed. This is the critical path \u2014 if the decision isn't "
            "made by June 2026, September is at risk."
        ],
        [
            "Macro dependencies",
            "Medium",
            f"{has_macros} workflows share macros. A macro used by 20 workflows "
            "blocks all 20 until it's migrated. A dependency map is needed."
        ],
        [
            "Ownership chasing",
            "Medium",
            f"{no_owner} workflows have no confirmed owner. People on leave, "
            "leavers, and unresponsive teams can burn 2\u20133 months of calendar time."
        ],
        [
            "Errored schedules",
            "Low",
            f"{errored_scheds} schedules are already in error. These may be "
            "broken today and could be decommissioned rather than migrated."
        ],
    ]

    add_styled_table(doc,
        ["Risk", "Severity", "Detail"],
        risks,
        col_widths=[4, 2, 11],
    )

    doc.add_paragraph()

    # ── Recommended Timeline ─────────────────────────────────────────────
    doc.add_heading("Recommended Timeline", level=1)

    doc.add_paragraph(
        "September 2026 is achievable if the Tier 4 product decision is made "
        "promptly and ownership resolution does not stall. Recommended phasing:"
    )

    timeline = [
        ["Phase 1: Cleanup & ownership", "Apr \u2013 May 2026", "6 weeks",
         "Decommission 566 inactive workflows. Resolve remaining ownership. "
         "Clean up 153 disabled schedules. Investigate 16 errored schedules."],
        ["Phase 2: Tier 1 migration", "Jun 2026", "4 weeks",
         f"Migrate {t1} simple ETL workflows to target platform."],
        ["Phase 3: Tier 2 migration", "Jun \u2013 Aug 2026", "8 weeks",
         f"Migrate {t2} orchestration workflows. Build macro dependency map first. "
         "This is the bulk of the work."],
        ["Phase 4: Tier 4 decision & build", "Apr \u2013 Aug 2026", "Parallel",
         f"Agree replacement for {t4} analytic apps. Must start immediately."],
        ["Phase 5: Cutover", "Sep 2026", "2 weeks",
         "Disable Alteryx Server. Monitor for breakage. Decommission."],
    ]

    add_styled_table(doc,
        ["Phase", "When", "Duration", "Scope"],
        timeline,
        col_widths=[4.5, 3, 2, 7.5],
    )

    doc.add_paragraph()

    # ── Next Steps ───────────────────────────────────────────────────────
    doc.add_heading("Recommended Next Steps", level=1)

    steps = [
        "Approve the migration scope: 161 active workflows, 566 to decommission.",
        "Make the Tier 4 product decision: what replaces the Alteryx analytic app interface?",
        f"Confirm ownership of the {no_owner} unresolved workflows via team leads." if no_owner > 0 else "Ownership is fully resolved \u2014 no action needed.",
        "Agree target platform per tier (ADF, Databricks, Python, or combination).",
        "Build the macro dependency map before starting Tier 2 migration.",
        "Prioritise the 43 scheduled workflows as the first migration cohort.",
        f"Investigate the {errored_scheds} errored schedules \u2014 some may already be decommission candidates.",
    ]

    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{step}", style="List Number")

    # ── Footer ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document generated from Alteryx Server API export data. April 2026.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(output_path)
    print(f"Briefing saved: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python Build-SLTBriefing.py <scoped_migration_report.xlsx> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.isfile(input_path):
        print(f"ERROR: {input_path} not found")
        sys.exit(1)

    output_path = sys.argv[2] if len(sys.argv) > 2 else os.path.join(
        os.path.dirname(input_path), "Alteryx_Migration_SLT_Briefing.docx"
    )

    print(f"Loading: {input_path}")
    data = load_data(input_path)
    print(f"  Workflows: {len(data['workflows'])}")
    print(f"  Schedules: {len(data['schedules'])}")
    print(f"  Job Activity: {len(data['jobs'])}")

    print("Generating briefing document...")
    build_document(data, output_path)


if __name__ == "__main__":
    main()
